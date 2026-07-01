#!/usr/bin/env python3
"""DOCX Report Generator - Word报告自动生成工具"""

import argparse
import sys
import json
import csv
import re
from pathlib import Path
from datetime import datetime

def fill_template(template_path, data_path, output=None):
    """模板填充"""
    from docx import Document
    import jinja2
    doc = Document(template_path)
    # 读取数据
    if data_path.endswith('.json'):
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    elif data_path.endswith('.csv'):
        import pandas as pd
        data = pd.read_csv(data_path).to_dict('records')
        if len(data) == 1:
            data = data[0]
    else:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    # 替换段落中的占位符
    for para in doc.paragraphs:
        for key, val in data.items():
            placeholder = '{{' + str(key) + '}}'
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(val))
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in data.items():
                        placeholder = '{{' + str(key) + '}}'
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(val))
    out = output or _auto_output(template_path, 'filled')
    doc.save(out)
    print(f"模板填充完成: {out}")

def create_report(title, config_path=None, output=None):
    """创建报告"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()
    # 如果有配置文件，按配置生成内容
    if config_path:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        for section in config.get('sections', []):
            stype = section.get('type', 'text')
            if stype == 'heading':
                doc.add_heading(section['content'], level=section.get('level', 1))
            elif stype == 'text':
                doc.add_paragraph(section['content'])
            elif stype == 'table':
                headers = section.get('headers', [])
                rows = section.get('rows', [])
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row):
                        table.rows[r_idx+1].cells[c_idx].text = str(val)
                doc.add_paragraph()
            elif stype == 'bullet':
                for item in section.get('items', []):
                    doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_heading('概述', level=1)
        doc.add_paragraph('请在此处填写报告概述...')
        doc.add_heading('详细分析', level=1)
        doc.add_paragraph('请在此处填写详细分析...')
        doc.add_heading('结论与建议', level=1)
        doc.add_paragraph('请在此处填写结论与建议...')
    out = output or f"{title}.docx"
    doc.save(out)
    print(f"报告已创建: {out}")

def mail_merge(template_path, data_path, output_dir=None):
    """批量邮件合并"""
    from docx import Document
    import pandas as pd
    out_dir = Path(output_dir or 'mail_merge_output')
    out_dir.mkdir(exist_ok=True)
    # 读取数据
    if data_path.endswith('.csv'):
        df = pd.read_csv(data_path, encoding='utf-8')
    elif data_path.endswith('.xlsx'):
        df = pd.read_excel(data_path)
    elif data_path.endswith('.json'):
        df = pd.read_json(data_path)
    else:
        df = pd.read_csv(data_path, encoding='utf-8')
    count = 0
    for idx, row in df.iterrows():
        doc = Document(template_path)
        data = row.to_dict()
        for para in doc.paragraphs:
            for key, val in data.items():
                placeholder = '{{' + str(key) + '}}'
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(val))
        for table in doc.tables:
            for t_row in table.rows:
                for cell in t_row.cells:
                    for para in cell.paragraphs:
                        for key, val in data.items():
                            placeholder = '{{' + str(key) + '}}'
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(val))
        filename = f"doc_{idx+1}.docx"
        doc.save(str(out_dir / filename))
        count += 1
    print(f"批量生成完成: {count}个文档 -> {out_dir}/")

def add_toc(filepath, output=None):
    """添加目录"""
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(filepath)
    # 在开头插入目录段落
    toc_para = doc.paragraphs[0].insert_paragraph_before('目录')
    toc_para.style = 'Heading 1'
    # 添加TOC域代码
    from docx.oxml import OxmlElement
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run = toc_para.add_run()
    run._element.append(fld_char_begin)
    run2 = toc_para.add_run()
    run2._element.append(instr_text)
    run3 = toc_para.add_run()
    run3._element.append(fld_char_end)
    out = output or _auto_output(filepath, 'with_toc')
    doc.save(out)
    print(f"目录已添加: {out} (在Word中打开后右键更新域)")

def to_pdf(filepath, output=None):
    """导出PDF"""
    out = output or Path(filepath).with_suffix('.pdf')
    os.system(f'libreoffice --headless --convert-to pdf "{filepath}" --outdir "{Path(out).parent}"')
    print(f"PDF已导出: {out}")

# --- Helpers ---

def _auto_output(filepath, suffix):
    p = Path(filepath)
    return str(p.parent / f"{p.stem}_{suffix}.docx")

def main():
    parser = argparse.ArgumentParser(description='DOCX Report Generator')
    sub = parser.add_subparsers(dest='command')

    p = sub.add_parser('fill-template')
    p.add_argument('template')
    p.add_argument('--data', required=True)
    p.add_argument('-o', '--output')

    p = sub.add_parser('create-report')
    p.add_argument('--title', required=True)
    p.add_argument('--config')
    p.add_argument('-o', '--output')

    p = sub.add_parser('mail-merge')
    p.add_argument('template')
    p.add_argument('--data', required=True)
    p.add_argument('-o', '--output')

    p = sub.add_parser('add-toc')
    p.add_argument('file')
    p.add_argument('-o', '--output')

    p = sub.add_parser('to-pdf')
    p.add_argument('file')
    p.add_argument('-o', '--output')

    args = parser.parse_args()
    if not args.command:
        parser.print_help()
        return
    cmds = {
        'fill-template': lambda: fill_template(args.template, args.data, args.output),
        'create-report': lambda: create_report(args.title, args.config, args.output),
        'mail-merge': lambda: mail_merge(args.template, args.data, args.output),
        'add-toc': lambda: add_toc(args.file, args.output),
        'to-pdf': lambda: to_pdf(args.file, args.output),
    }
    cmds[args.command]()

if __name__ == '__main__':
    main()
