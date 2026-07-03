"""测试专用 Word 文档适配器。"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from _skill.models import SkillDefinition


class TestWordDocumentAdapter:
    """最小 Word 产物生成 fake，用于端到端 artifact 测试。"""

    name = "TestWordDocument"

    def execute(self, prompt: str, *, skill: SkillDefinition, context: dict[str, Any]) -> dict[str, Any]:
        """根据测试字段生成 docx 文件并返回产物路径。"""
        from docx import Document

        fields = context.get("fields", {})
        output_path = Path(fields.get("output_path") or fields.get("filename") or "skill-output.docx")
        if not output_path.is_absolute():
            output_path = Path.cwd() / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(str(fields.get("title") or skill.name), level=1)
        document.add_paragraph(str(fields.get("content") or prompt[:1000]))
        document.save(output_path)
        return {"output_path": str(output_path), "skill_name": skill.name}
