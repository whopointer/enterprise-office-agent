"""适配器契约自动化测试。"""

from __future__ import annotations

from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path
from threading import Thread
import json

from docx import Document

from _skill import FileSkillDiscovery
from _skill.models import CallableSkillAdapter
from adapters.skill_adapters import (
    LangChainSkillAdapter,
    OpenAICompatibleSkillAdapter,
    SpringAIHttpAdapter,
)
from core.executor import SkillExecutor
from tests.fakes import TestWordDocumentAdapter
from tests.skill_fixtures import build_pipeline_test_skills


class _Message:
    def __init__(self, content: str):
        self.content = content


class _Choice:
    def __init__(self, content: str):
        self.message = _Message(content)


class _Usage:
    def __init__(self, prompt_tokens: int, completion_tokens: int, total_tokens: int):
        self.prompt_tokens = prompt_tokens
        self.completion_tokens = completion_tokens
        self.total_tokens = total_tokens


class _Response:
    def __init__(self, content: str, usage: _Usage | None = None):
        self.choices = [_Choice(content)]
        self.usage = usage


class _FakeCompletions:
    def __init__(self, content: str | None = "adapter-ok", error: Exception | None = None, usage: _Usage | None = None):
        self.content = content
        self.error = error
        self.usage = usage
        self.calls: list[dict] = []

    def create(self, **kwargs):
        self.calls.append(kwargs)
        if self.error:
            raise self.error
        return _Response(self.content, self.usage)


class _FakeChat:
    def __init__(self, completions: _FakeCompletions):
        self.completions = completions


class _FakeOpenAIClient:
    def __init__(self, content: str | None = "adapter-ok", error: Exception | None = None, usage: _Usage | None = None):
        self.completions = _FakeCompletions(content, error, usage)
        self.chat = _FakeChat(self.completions)


def _simple_skill(tmp_path: Path):
    index = FileSkillDiscovery(build_pipeline_test_skills(tmp_path)).discover()
    return index, index.get("simple-echo")


def test_openai_compatible_adapter_contract_with_fake_client(tmp_path: Path) -> None:
    """OpenAI adapter 应按 Chat Completions 协议发送 prompt 并返回文本。"""
    index, skill = _simple_skill(tmp_path)
    client = _FakeOpenAIClient()
    adapter = OpenAICompatibleSkillAdapter(model="fake-model", client=client)

    result = SkillExecutor(index, adapter).execute(skill, user_query="ping")

    assert result.metrics.execution_success is True
    assert result.output == "adapter-ok"
    assert client.completions.calls
    assert client.completions.calls[0]["model"] == "fake-model"
    assert client.completions.calls[0]["messages"][1]["content"] == result.context.prompt
    assert result.metrics.token_metrics.source == "estimated"


def test_openai_compatible_adapter_uses_actual_token_usage(tmp_path: Path) -> None:
    """供应商返回 usage 时，执行器应优先记录真实 token 数。"""
    index, skill = _simple_skill(tmp_path)
    client = _FakeOpenAIClient(usage=_Usage(prompt_tokens=11, completion_tokens=7, total_tokens=18))
    adapter = OpenAICompatibleSkillAdapter(model="fake-model", client=client)

    result = SkillExecutor(index, adapter).execute(skill, user_query="ping")

    assert result.metrics.execution_success is True
    assert result.metrics.token_metrics.input_tokens == 11
    assert result.metrics.token_metrics.output_tokens == 7
    assert result.metrics.token_metrics.total_tokens == 18
    assert result.metrics.token_metrics.source == "actual"


def test_openai_compatible_adapter_empty_content_returns_empty_string(tmp_path: Path) -> None:
    """OpenAI adapter 遇到空 content 时应返回空字符串而不是 None。"""
    index, skill = _simple_skill(tmp_path)
    client = _FakeOpenAIClient(content=None)

    result = SkillExecutor(index, OpenAICompatibleSkillAdapter(model="fake-model", client=client)).execute(skill, user_query="ping")

    assert result.metrics.execution_success is True
    assert result.output == ""


def test_openai_compatible_adapter_exception_is_captured(tmp_path: Path) -> None:
    """OpenAI adapter 抛异常时由执行器兜底为失败结果。"""
    index, skill = _simple_skill(tmp_path)
    client = _FakeOpenAIClient(error=RuntimeError("provider down"))

    result = SkillExecutor(index, OpenAICompatibleSkillAdapter(model="fake-model", client=client)).execute(skill, user_query="ping")

    assert result.metrics.execution_success is False
    assert result.output["error"] == "provider down"


def test_langchain_adapter_exposes_usage_for_executor(tmp_path: Path) -> None:
    """LangChain 返回 token_usage 时，执行器应记录真实 token。"""
    index, skill = _simple_skill(tmp_path)

    class Runnable:
        def invoke(self, _payload):
            return {
                "content": "ok",
                "llm_output": {
                    "token_usage": {
                        "prompt_tokens": 21,
                        "completion_tokens": 9,
                        "total_tokens": 30,
                    }
                },
            }

    result = SkillExecutor(index, LangChainSkillAdapter(Runnable())).execute(skill, user_query="ping")

    assert result.metrics.execution_success is True
    assert result.metrics.token_metrics.input_tokens == 21
    assert result.metrics.token_metrics.output_tokens == 9
    assert result.metrics.token_metrics.total_tokens == 30
    assert result.metrics.token_metrics.source == "actual"


def test_spring_ai_adapter_exposes_usage_for_executor(tmp_path: Path) -> None:
    """SpringAI HTTP 返回 usage 时，执行器应记录真实 token。"""
    index, skill = _simple_skill(tmp_path)

    class Handler(BaseHTTPRequestHandler):
        def do_POST(self):
            self.rfile.read(int(self.headers["Content-Length"]))
            body = json.dumps({
                "content": "ok",
                "usage": {"prompt_tokens": 13, "completion_tokens": 4, "total_tokens": 17},
            }).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        def log_message(self, *_args):
            return

    server = HTTPServer(("127.0.0.1", 0), Handler)
    thread = Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        endpoint = f"http://127.0.0.1:{server.server_port}"
        result = SkillExecutor(index, SpringAIHttpAdapter(endpoint)).execute(skill, user_query="ping")
    finally:
        server.shutdown()
        thread.join(timeout=2)

    assert result.metrics.execution_success is True
    assert result.metrics.token_metrics.input_tokens == 13
    assert result.metrics.token_metrics.output_tokens == 4
    assert result.metrics.token_metrics.total_tokens == 17
    assert result.metrics.token_metrics.source == "actual"


def test_word_document_adapter_contract_reads_generated_docx(tmp_path: Path) -> None:
    """TestWordDocument adapter 不只验证文件存在，还要读回标题和正文。"""
    index, skill = _simple_skill(tmp_path)
    output_path = tmp_path / "contract.docx"

    result = SkillExecutor(index, TestWordDocumentAdapter()).execute(
        skill,
        user_query="生成文档",
        fields={"output_path": str(output_path), "title": "合同测试", "content": "正文内容"},
    )
    document = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert result.metrics.execution_success is True
    assert output_path.exists()
    assert "合同测试" in text
    assert "正文内容" in text


def test_word_document_adapter_overwrites_existing_file(tmp_path: Path) -> None:
    """TestWordDocument adapter 重复写同一路径时应生成新的可读内容。"""
    index, skill = _simple_skill(tmp_path)
    output_path = tmp_path / "overwrite.docx"
    output_path.write_text("old", encoding="utf-8")

    result = SkillExecutor(index, TestWordDocumentAdapter()).execute(
        skill,
        user_query="生成文档",
        fields={"output_path": str(output_path), "title": "新标题", "content": "新内容"},
    )
    text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

    assert result.metrics.execution_success is True
    assert "新标题" in text
    assert "新内容" in text


def test_word_document_adapter_relative_output_path(tmp_path: Path, monkeypatch) -> None:
    """相对 output_path 应相对当前工作目录生成。"""
    monkeypatch.chdir(tmp_path)
    index, skill = _simple_skill(tmp_path)

    result = SkillExecutor(index, TestWordDocumentAdapter()).execute(
        skill,
        user_query="生成文档",
        fields={"output_path": "nested/relative.docx", "title": "相对路径"},
    )

    assert result.metrics.execution_success is True
    assert (tmp_path / "nested" / "relative.docx").exists()


def test_word_document_adapter_defaults_content_from_prompt(tmp_path: Path) -> None:
    """未提供 content 时应把 prompt 摘要写入文档。"""
    index, skill = _simple_skill(tmp_path)
    output_path = tmp_path / "prompt-content.docx"

    result = SkillExecutor(index, TestWordDocumentAdapter()).execute(
        skill,
        user_query="生成文档",
        fields={"output_path": str(output_path), "title": "Prompt 摘要"},
    )
    text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

    assert result.metrics.execution_success is True
    assert "# Skill: simple-echo" in text


class _JsonHandler(BaseHTTPRequestHandler):
    def do_POST(self):  # noqa: N802
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length).decode("utf-8")
        payload = json.loads(body)
        raw = json.dumps({"received_skill": payload["skill"], "has_prompt": bool(payload["prompt"])}).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def log_message(self, _format, *_args):
        return


def test_springai_adapter_contract_with_local_http_server(tmp_path: Path) -> None:
    """SpringAI adapter 应能向 HTTP 服务 POST prompt/context 并解析 JSON。"""
    server = HTTPServer(("127.0.0.1", 0), _JsonHandler)
    thread = Thread(target=server.serve_forever, daemon=True)
    thread.start()

    try:
        index, skill = _simple_skill(tmp_path)
        endpoint = f"http://127.0.0.1:{server.server_port}/skill"
        result = SkillExecutor(index, SpringAIHttpAdapter(endpoint, timeout=3)).execute(skill, user_query="ping")
    finally:
        server.shutdown()
        thread.join(timeout=3)

    assert result.metrics.execution_success is True
    assert result.output == {"received_skill": "simple-echo", "has_prompt": True}


class _TextHandler(BaseHTTPRequestHandler):
    def do_POST(self):  # noqa: N802
        raw = b"plain-text-ok"
        self.send_response(200)
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def log_message(self, _format, *_args):
        return


class _ErrorHandler(BaseHTTPRequestHandler):
    def do_POST(self):  # noqa: N802
        raw = b"{\"error\":\"bad request\"}"
        self.send_response(500)
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def log_message(self, _format, *_args):
        return


def _run_server(handler):
    server = HTTPServer(("127.0.0.1", 0), handler)
    thread = Thread(target=server.serve_forever, daemon=True)
    thread.start()
    return server, thread


def test_springai_adapter_returns_plain_text_for_non_json_response(tmp_path: Path) -> None:
    """SpringAI adapter 对 200 非 JSON 响应应返回原始文本。"""
    server, thread = _run_server(_TextHandler)
    try:
        index, skill = _simple_skill(tmp_path)
        result = SkillExecutor(index, SpringAIHttpAdapter(f"http://127.0.0.1:{server.server_port}/x", timeout=3)).execute(skill, user_query="ping")
    finally:
        server.shutdown()
        thread.join(timeout=3)

    assert result.metrics.execution_success is True
    assert result.output == "plain-text-ok"


def test_springai_adapter_http_500_is_captured(tmp_path: Path) -> None:
    """SpringAI adapter 对 HTTP 500 应返回执行失败而不是崩溃。"""
    server, thread = _run_server(_ErrorHandler)
    try:
        index, skill = _simple_skill(tmp_path)
        result = SkillExecutor(index, SpringAIHttpAdapter(f"http://127.0.0.1:{server.server_port}/x", timeout=3)).execute(skill, user_query="ping")
    finally:
        server.shutdown()
        thread.join(timeout=3)

    assert result.metrics.execution_success is False
    assert "SpringAI HTTP 500" in result.output["error"]


class _TypeErrorRunner:
    def invoke(self, payload):
        if isinstance(payload, dict):
            raise TypeError("dict unsupported")
        return f"fallback:{payload[:6]}"


def test_langchain_adapter_falls_back_to_prompt_on_type_error(tmp_path: Path) -> None:
    """LangChain invoke(payload) TypeError 时应退回 invoke(prompt)。"""
    from adapters.skill_adapters import LangChainSkillAdapter

    index, skill = _simple_skill(tmp_path)
    result = SkillExecutor(index, LangChainSkillAdapter(_TypeErrorRunner())).execute(skill, user_query="hello")

    assert result.metrics.execution_success is True
    assert str(result.output).startswith("fallback:")


def test_callable_adapter_can_return_structured_payload(tmp_path: Path) -> None:
    """Callable adapter 应允许返回 dict 等结构化对象。"""
    index, skill = _simple_skill(tmp_path)
    adapter = CallableSkillAdapter("structured", lambda _prompt, current_skill, context: {
        "skill": current_skill.name,
        "query": context["user_query"],
        "fields": context["fields"],
    })

    result = SkillExecutor(index, adapter).execute(skill, user_query="ping", fields={"filename": "x.docx"})

    assert result.metrics.execution_success is True
    assert result.output == {
        "skill": "simple-echo",
        "query": "ping",
        "fields": {"filename": "x.docx"},
    }
