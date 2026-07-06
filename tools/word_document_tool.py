"""Word 文档生成工具。"""

from __future__ import annotations

from pathlib import Path
from typing import Any


class WordDocumentTool:
    """基于 python-docx 生成最小 Word 文档。"""

    name = "WordDocumentTool"

    def run(self, *, fields: dict[str, Any], content: str = "", default_title: str = "Document") -> dict[str, Any]:
        """生成 docx 文件并返回产物路径。"""
        from docx import Document

        output_path = Path(fields.get("output_path") or fields.get("filename") or "skill-output.docx")
        if not output_path.is_absolute():
            output_path = Path.cwd() / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(str(fields.get("title") or default_title), level=1)
        document.add_paragraph(str(fields.get("content") or content[:1000]))
        document.save(output_path)
        return {"output_path": str(output_path)}
