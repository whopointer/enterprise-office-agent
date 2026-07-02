"""Skill 执行适配器。"""

from __future__ import annotations

from pathlib import Path
from typing import Any
import json
import urllib.error
import urllib.request

from openai import OpenAI

from _skill.models import SkillDefinition
from llm.skill_router import _build_openai_client, load_skill_env


class OpenAICompatibleSkillAdapter:
    """把已组装的 skill prompt 交给 OpenAI-compatible chat API 执行。"""

    name = "OpenAICompatible"

    def __init__(self, *, model: str | None = None, client: OpenAI | None = None) -> None:
        load_skill_env()
        import os

        self.model = model or os.environ.get("MODEL") or os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
        self.client = client or _build_openai_client()

    def execute(self, prompt: str, *, skill: SkillDefinition, context: dict[str, Any]) -> str:
        """调用大模型执行 skill prompt。"""
        response = self.client.chat.completions.create(
            model=self.model,
            temperature=0.2,
            messages=[
                {"role": "system", "content": f"你正在执行 skill: {skill.name}。请严格遵循上下文。"},
                {"role": "user", "content": prompt},
            ],
        )
        return response.choices[0].message.content or ""


class LangChainSkillAdapter:
    """包装 LangChain Runnable / Chain，适配 SkillAdapter 协议。"""

    name = "LangChain"

    def __init__(self, runnable: Any) -> None:
        self.runnable = runnable

    def execute(self, prompt: str, *, skill: SkillDefinition, context: dict[str, Any]) -> Any:
        """调用 LangChain 的 invoke 接口。"""
        payload = {"input": prompt, "skill": skill.name, "context": context}
        if hasattr(self.runnable, "invoke"):
            try:
                return self.runnable.invoke(payload)
            except TypeError:
                return self.runnable.invoke(prompt)
        if callable(self.runnable):
            return self.runnable(prompt)
        raise TypeError("LangChainSkillAdapter 需要 runnable.invoke 或 callable")


class SpringAIHttpAdapter:
    """通过 HTTP 调用 SpringAI 服务。"""

    name = "SpringAI"

    def __init__(self, endpoint: str, *, timeout: float = 30.0) -> None:
        self.endpoint = endpoint
        self.timeout = timeout

    def execute(self, prompt: str, *, skill: SkillDefinition, context: dict[str, Any]) -> Any:
        """向 SpringAI HTTP endpoint POST prompt 和上下文。"""
        body = json.dumps(
            {"prompt": prompt, "skill": skill.name, "context": context},
            ensure_ascii=False,
            default=str,
        ).encode("utf-8")
        request = urllib.request.Request(
            self.endpoint,
            data=body,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=self.timeout) as response:
                raw = response.read().decode("utf-8")
        except urllib.error.HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"SpringAI HTTP {exc.code}: {raw}") from exc
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            return raw


class WordDocumentSkillAdapter:
    """最小 Word 产物生成 adapter，用于端到端 artifact 测试。"""

    name = "WordDocument"

    def execute(self, prompt: str, *, skill: SkillDefinition, context: dict[str, Any]) -> dict[str, Any]:
        """根据字段生成一个 docx 文件并返回产物路径。"""
        from docx import Document

        fields = context.get("fields", {})
        output_path = Path(fields.get("output_path") or fields.get("filename") or "skill-output.docx")
        if not output_path.is_absolute():
            output_path = Path.cwd() / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading(str(fields.get("title") or skill.name), level=1)
        document.add_paragraph(str(fields.get("content") or prompt[:1000]))
        document.save(output_path)
        return {"output_path": str(output_path), "skill_name": skill.name}
